#! /usr/bin/env python3
import os
import csv
import sys
import argparse
from docx import Document

def process_docx_file(input_path):
    """处理单个DOCX文件"""
    try:
        doc = Document(input_path)
    except Exception as e:
        print(f"错误：无法读取文件 '{input_path}' ({str(e)})", file=sys.stderr)
        return False

    if not doc.tables:
        print(f"提示：{os.path.basename(input_path)} 中没有表格，已跳过", file=sys.stderr)
        return False

    # 生成输出路径
    output_dir = os.path.dirname(input_path)
    base_name = os.path.splitext(os.path.basename(input_path))[0]
    output_path = os.path.join(output_dir, f"{base_name}.csv")

    # 提取第一个表格的表头
    header = []
    first_table = doc.tables[0]
    for cell in first_table.rows[0].cells:
        header.append(cell.text.strip())

    # 收集所有数据行
    data_rows = []
    for table_idx, table in enumerate(doc.tables):
        start_row = 0
        if table_idx == 0:
            start_row = 1  # 跳过第一个表格的表头
        else:
            # 检查后续表格的首行是否为重复表头
            current_first_row = [cell.text.strip() for cell in table.rows[0].cells]
            if current_first_row == header:
                start_row = 1

        # 提取数据行
        for row in table.rows[start_row:]:
            row_data = [str(cell.text.strip()) for cell in row.cells]
            if len(row_data) != len(header):
                print(f"警告：{os.path.basename(input_path)} 中跳过列数不一致的行", file=sys.stderr)
                continue
            data_rows.append(row_data)

    # 写入CSV文件
    with open(output_path, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f, quoting=csv.QUOTE_MINIMAL)
        writer.writerow(header)
        writer.writerows(data_rows)
    
    print(f"成功生成：{output_path}")
    return True

def process_directory(input_dir):
    """处理整个目录"""
    for filename in os.listdir(input_dir):
        if filename.lower().endswith('.docx'):
            file_path = os.path.join(input_dir, filename)
            process_docx_file(file_path)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="批量转换DOCX表格为CSV文件")
    parser.add_argument("input_dir", help="包含DOCX文件的目录路径")
    args = parser.parse_args()
    
    if not os.path.isdir(args.input_dir):
        print(f"错误：{args.input_dir} 不是有效目录", file=sys.stderr)
        sys.exit(1)
        
    process_directory(args.input_dir)
