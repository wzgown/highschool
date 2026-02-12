#! /usr/bin/env python3
from docx import Document
import os
import glob
import re

def convert_docx_to_markdown(docx_path, output_path):
    """将单个DOCX文件转换为Markdown"""
    try:
        doc = Document(docx_path)
        markdown = []
        seen_tables = set()  # 用于记录已处理过的表格内容
        
        # 先处理所有段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            # 标题识别
            if re.match(r'^[一二三四五六七八九十]+、', text):
                markdown.append(f"\n### {text}\n")
            elif any(char in text for char in ["●", "•", ">", "√"]):
                markdown.append(f"- {text.replace('>', '').strip()}")
            else:
                markdown.append(text)
        
        # 再处理所有表格
        for table in doc.tables:
            try:
                df = []
                for row in table.rows:
                    row_data = [cell.text.strip().replace('\n', '<br>') for cell in row.cells]
                    df.append(row_data)
                
                if df:
                    # 生成表格内容的唯一标识
                    table_content = str(df)
                    if table_content in seen_tables:
                        continue  # 跳过重复的表格
                    seen_tables.add(table_content)
                    
                    headers = df[0]
                    markdown_table = [
                        "\n| " + " | ".join(headers) + " |",
                        "|" + "|".join(["---"] * len(headers)) + "|"
                    ]
                    for row in df[1:]:
                        markdown_table.append("| " + " | ".join(row) + " |")
                    markdown.append("\n".join(markdown_table) + "\n")
                    
            except Exception as e:
                print(f"警告：处理表格时发生错误（{str(e)}），已跳过")
                continue

        # 后处理清理
        cleaned_md = []
        for line in markdown:
            line = re.sub(r'\s{2,}', ' ', line)  # 合并多余空格
            line = line.replace(" ", " ")        # 替换全角空格
            if line.strip():
                cleaned_md.append(line)

        # 写入文件
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("\n".join(cleaned_md))
        return True
    
    except Exception as e:
        print(f"转换失败: {docx_path}\n错误信息: {str(e)}")
        return False

def batch_convert(directory):
    """批量转换目录下的所有DOCX文件"""
    docx_files = glob.glob(os.path.join(directory, "*.docx"))
    total = len(docx_files)
    success = 0
    
    for docx_path in docx_files:
        # 生成输出路径
        base = os.path.splitext(docx_path)[0]
        output_path = f"{base}.md"
        
        # 跳过已存在的文件
        if os.path.exists(output_path):
            print(f"已跳过（文件已存在）: {output_path}")
            continue
            
        # 执行转换
        if convert_docx_to_markdown(docx_path, output_path):
            print(f"转换成功: {output_path}")
            success += 1
    
    print(f"\n转换完成！成功: {success}/{total} 文件")

if __name__ == "__main__":
    target_directory = input("请输入要转换的目录路径: ").strip()
    if os.path.isdir(target_directory):
        batch_convert(target_directory)
    else:
        print("错误：输入的路径不是有效目录！")